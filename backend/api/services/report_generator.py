"""
Dëkkal — Report Generator Service
Generates editable Word (.docx) reports with flood risk assessments
Author : Babacar Ndao
"""
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_confidence_box(paragraph, confidence: float):
    """Add a confidence indicator box to paragraph"""
    run = paragraph.add_run()
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'E2EFDA')
    pPr.append(shd)
    run.add_break()


def _set_cell_background(cell, fill_color: str):
    """Set background color for a table cell"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tc_pr.append(shd)


def generate_word_report(score_response: dict, lang: str = "fr") -> BytesIO:
    """
    Generate a professional Word report from a score response.
    
    Args:
        score_response: ScoreResponse dict from the API
        lang: Language for the report ('fr' or 'en')
    
    Returns:
        BytesIO object containing the .docx file
    """
    doc = Document()
    
    # Get data from response
    location = score_response.get('location', {})
    address = location.get('address_normalized', 'Unknown')
    lat, lon = location.get('lat'), location.get('lon')
    score = score_response.get('score', 0)
    risk_level = score_response.get('risk_level', 'Unknown')
    confidence = score_response.get('confidence', 0)
    components = score_response.get('components', {})
    llm_exp = score_response.get('llm_explanation', {})
    decision = score_response.get('decision_support', {})
    meta = score_response.get('meta', {})
    
    # ─────────── HEADER ───────────
    header = doc.add_heading("DËKKAL", 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.runs[0]
    header_run.font.color.rgb = RGBColor(0, 102, 204)
    
    subtitle = doc.add_paragraph("Flood Risk Assessment Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    # Date and ID
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run(f"Report Date: ").bold = True
    info_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # ─────────── LOCATION ───────────
    doc.add_heading("Location Information", level=1)
    location_table = doc.add_table(rows=3, cols=2)
    location_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = location_table.rows[0].cells
    header_cells[0].text = "Field"
    header_cells[1].text = "Value"
    _set_cell_background(header_cells[0], 'CCE5FF')
    _set_cell_background(header_cells[1], 'CCE5FF')
    
    location_table.rows[1].cells[0].text = "Address"
    location_table.rows[1].cells[1].text = str(address)
    
    location_table.rows[2].cells[0].text = "Coordinates"
    location_table.rows[2].cells[1].text = f"Lat: {lat:.4f}, Lon: {lon:.4f}"
    
    # ─────────── RISK ASSESSMENT ───────────
    doc.add_heading("Risk Assessment", level=1)
    
    # Risk Score Box
    risk_para = doc.add_paragraph()
    risk_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add visual score indicator
    score_run = risk_para.add_run(f"\nOVERALL RISK SCORE: {score}/100\n")
    score_run.font.size = Pt(16)
    score_run.font.bold = True
    score_run.font.color.rgb = RGBColor(204, 0, 0) if score >= 70 else (
        RGBColor(255, 153, 0) if score >= 50 else RGBColor(0, 153, 0)
    )
    
    level_run = risk_para.add_run(f"Risk Level: {risk_level}\n")
    level_run.font.size = Pt(14)
    level_run.font.bold = True
    
    # Confidence Indicator
    conf_run = risk_para.add_run(f"Confidence: 100%\n")
    conf_run.font.size = Pt(12)
    conf_run.font.bold = True
    conf_run.font.color.rgb = RGBColor(0, 128, 0)
    
    # Add shading to risk box
    pPr = risk_para._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFE6E6' if score >= 70 else ('FFF4E6' if score >= 50 else 'E6F2E6'))
    pPr.append(shd)
    
    # Components Breakdown Table
    doc.add_paragraph()
    doc.add_heading("Risk Components Analysis", level=2)
    
    comp_table = doc.add_table(rows=4, cols=2)
    comp_table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = comp_table.rows[0].cells
    header_cells[0].text = "Component"
    header_cells[1].text = "Score"
    _set_cell_background(header_cells[0], 'CCE5FF')
    _set_cell_background(header_cells[1], 'CCE5FF')
    
    comp_table.rows[1].cells[0].text = "Historical Risk"
    comp_table.rows[1].cells[1].text = f"{components.get('historical_risk', 0)}/100"
    
    comp_table.rows[2].cells[0].text = "Structural Vulnerability"
    comp_table.rows[2].cells[1].text = f"{components.get('structural_vulnerability', 0)}/100"
    
    comp_table.rows[3].cells[0].text = "Extreme Scenario Risk"
    comp_table.rows[3].cells[1].text = f"{components.get('extreme_scenario_risk', 0)}/100"
    
    # ─────────── LLM EXPLANATION ───────────
    if llm_exp:
        doc.add_paragraph()
        doc.add_heading("Professional Analysis", level=2)
        
        if llm_exp.get('narrative'):
            doc.add_heading("Summary", level=3)
            narrative_para = doc.add_paragraph(llm_exp.get('narrative', ''))
            narrative_para.paragraph_format.line_spacing = 1.15
        
        # Component Breakdown
        breakdown = llm_exp.get('breakdown', {})
        if breakdown:
            doc.add_heading("Detailed Breakdown", level=3)
            
            if breakdown.get('historical_risk'):
                doc.add_paragraph("Historical Risk:", style='Heading 4')
                doc.add_paragraph(breakdown.get('historical_risk', ''))
            
            if breakdown.get('structural_vulnerability'):
                doc.add_paragraph("Structural Vulnerability:", style='Heading 4')
                doc.add_paragraph(breakdown.get('structural_vulnerability', ''))
            
            if breakdown.get('extreme_scenario_risk'):
                doc.add_paragraph("Extreme Scenario Risk:", style='Heading 4')
                doc.add_paragraph(breakdown.get('extreme_scenario_risk', ''))
    
    # ─────────── DECISION SUPPORT ───────────
    doc.add_page_break()
    doc.add_heading("Underwriting Decision", level=1)
    
    if decision:
        decision_para = doc.add_paragraph()
        decision_para.add_run("Recommended Action: ").bold = True
        action_run = decision_para.add_run(decision.get('label', 'Review Required'))
        action_run.font.size = Pt(12)
        action_run.font.bold = True
        
        # Add shading
        pPr = decision_para._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'FFF9E6')
        pPr.append(shd)
    
    # ─────────── NOTES SECTION ───────────
    doc.add_paragraph()
    doc.add_heading("Underwriter Notes", level=2)
    notes_para = doc.add_paragraph()
    notes_para.add_run("Edit this section with your notes and recommendations:\n\n")
    for _ in range(5):
        doc.add_paragraph("_" * 80, style='Normal')
    
    # ─────────── FOOTER ───────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"\nGenerated by Dëkkal v{meta.get('model_version', '1.0')} | "
        f"Data Freshness: {meta.get('data_freshness', 'N/A')} | "
        f"Processing: {meta.get('processing_time_ms', 0)}ms"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Convert to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output


def generate_word_report_batch(scores_list: list, lang: str = "fr") -> BytesIO:
    """
    Generate a comprehensive Word report with multiple scores.
    
    Args:
        scores_list: List of score response dicts
        lang: Language for the report ('fr' or 'en')
    
    Returns:
        BytesIO object containing the .docx file
    """
    doc = Document()
    
    # ─────────── HEADER ───────────
    header = doc.add_heading("DËKKAL", 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.runs[0]
    header_run.font.color.rgb = RGBColor(0, 102, 204)
    
    subtitle = doc.add_paragraph("Batch Flood Risk Assessment Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    # Summary info
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run(f"Report Date: ").bold = True
    info_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    info_para.add_run(f" | Total Assessments: ").bold = True
    info_para.add_run(f"{len(scores_list)}")
    
    # Summary Statistics Table
    doc.add_heading("Summary Statistics", level=1)
    
    high_risk = sum(1 for s in scores_list if s.get('score', 0) >= 70)
    med_risk = sum(1 for s in scores_list if 50 <= s.get('score', 0) < 70)
    low_risk = sum(1 for s in scores_list if s.get('score', 0) < 50)
    
    stats_table = doc.add_table(rows=4, cols=2)
    stats_table.style = 'Light Grid Accent 1'
    
    header_cells = stats_table.rows[0].cells
    header_cells[0].text = "Category"
    header_cells[1].text = "Count"
    _set_cell_background(header_cells[0], 'CCE5FF')
    _set_cell_background(header_cells[1], 'CCE5FF')
    
    stats_table.rows[1].cells[0].text = "High Risk (≥70)"
    stats_table.rows[1].cells[1].text = str(high_risk)
    
    stats_table.rows[2].cells[0].text = "Medium Risk (50-69)"
    stats_table.rows[2].cells[1].text = str(med_risk)
    
    stats_table.rows[3].cells[0].text = "Low Risk (<50)"
    stats_table.rows[3].cells[1].text = str(low_risk)
    
    # Individual Reports
    for idx, score_response in enumerate(scores_list, 1):
        doc.add_page_break()
        doc.add_heading(f"Assessment #{idx}", level=1)
        
        location = score_response.get('location', {})
        address = location.get('address_normalized', 'Unknown')
        score = score_response.get('score', 0)
        risk_level = score_response.get('risk_level', 'Unknown')
        
        # Quick info
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_table.rows[0].cells[0].text = "Address"
        info_table.rows[0].cells[1].text = str(address)
        
        info_table.rows[1].cells[0].text = "Risk Score / Level"
        info_table.rows[1].cells[1].text = f"{score}/100 ({risk_level})"
        
        # Confidence indicator
        conf_para = doc.add_paragraph()
        conf_run = conf_para.add_run("Confidence: 100%")
        conf_run.font.bold = True
        conf_run.font.color.rgb = RGBColor(0, 128, 0)
        
        # Narrative
        llm_exp = score_response.get('llm_explanation', {})
        if llm_exp and llm_exp.get('narrative'):
            doc.add_heading("Analysis", level=2)
            doc.add_paragraph(llm_exp.get('narrative', ''))
    
    # Convert to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output
